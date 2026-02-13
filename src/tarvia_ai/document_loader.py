from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List

from docx import Document as DocxDocument
from pypdf import PdfReader


SUPPORTED_SUFFIXES = {".txt", ".md", ".pdf", ".docx"}


@dataclass(frozen=True)
class LoadedDocument:
    doc_id: str
    source: str
    text: str


def load_documents(paths: Iterable[Path]) -> List[LoadedDocument]:
    documents: List[LoadedDocument] = []

    for raw_path in paths:
        path = raw_path.expanduser().resolve()
        if not path.exists():
            continue
        if path.suffix.lower() not in SUPPORTED_SUFFIXES:
            continue

        text = _read_text(path)
        if not text.strip():
            continue

        documents.append(
            LoadedDocument(
                doc_id=_safe_doc_id(path),
                source=str(path),
                text=text,
            )
        )
    return documents


def _read_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx":
        doc = DocxDocument(str(path))
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)
    return ""


def _safe_doc_id(path: Path) -> str:
    return path.stem.replace(" ", "_").replace(".", "_")
